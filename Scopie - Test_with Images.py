'''Commands for module installation
--pip install docxtpl
pip install openpyxl
pip install python-docx
pip install pandas
--pip install xlrd
'''
import time
import sys
import os
import docx
from docx.shared import Inches
from openpyxl import load_workbook
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_FILE_NAME = "InputSheet_Excel_1.xlsx"  #Input File excel
OUTPUT_FILE_NAME = "NEOMU" + "_OUTPUT.docx"  #Output File word
DATA_SHEET = "InputData"
TEMPLATE_SHEET = "Template"
FIELD_SELECTION_SHEET = "Topics2Pick"
EXCEL_DELIMITER = "-=-"
IMAGE_WIDTH = 4.9
KEY_H1_MAPPING = "H1"
KEY_H2_MAPPING = "H2"
KEY_H3_MAPPING = "H3"
KEY_H4_MAPPING = "H4"
KEY_H5_MAPPING = "H5"
KEY_H6_MAPPING = "H6"

# Styles

H1_STYLE = "Heading 1"
H2_STYLE = "Heading 2"
H3_STYLE = "Heading 3"
H4_STYLE = "Heading 4"
H5_STYLE = "Heading 5"
H6_STYLE = "Heading 6"
PARA_STYLE = "Body Text"
# Variables to remember last updated headings, so that it does not repeat printing them on the word doc
last_h1_updated = last_h2_updated = last_h3_updated = last_h4_updated = last_h5_updated = last_h6_updated = ""


# Function to print content,styles and Images excluding tables
def add_paragraph(doc, content, style, img_path=""):
    p = doc.add_paragraph()
    if len(str(style)) > 0:
        p.style = doc.styles[style]
    run = p.add_run(content)
    if img_path is not None and len(img_path) > 0:
        run.add_picture(img_path, width=Inches(IMAGE_WIDTH))  #, height=Inches(.7))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(content)


# Function to read the input excel sheet and return the data to process the document
def read_input_sheet_data():
    file_path = "./" + INPUT_FILE_NAME
    file = None
    data = []
    selection_map = dict()

    if os.path.exists(file_path) and os.path.isfile(file_path):
        try:
            file = INPUT_FILE_NAME
            wb_obj = load_workbook(filename=file)
            wsheet = wb_obj[DATA_SHEET]
            for cell in wsheet.iter_rows(min_row=2, values_only=True):
                res = [str(i or '') for i in cell]
                data.append(EXCEL_DELIMITER.join(res))

            wsheet = None
            wsheet = wb_obj[FIELD_SELECTION_SHEET]

            for col in wsheet.iter_cols(values_only=True):
                res = [x for x in col if x is not None]
                if res is not None and len(res) > 0:
                    key = res[0].strip()
                    res.pop(0)
                    data_set = set()

                    for val in res:
                        data_set.add(val.strip())
                    selection_map[key] = data_set
                    print(selection_map)

            template_sheet = wb_obj[TEMPLATE_SHEET]
            if template_sheet.max_row == 1:
                template_filepath = template_sheet.cell(row=1, column=1).value
            else:
                template_filepath = None
            return data, selection_map, template_filepath
        except Exception as err:
            print("ERROR: Not able to read input sheet data. Aborting job!!!\n" % (INPUT_FILE_NAME, str(err)))
            sys.exit()
    else:
        print("ERROR: Input sheet data file (%s) does not exists. Aborting job!!!" % INPUT_FILE_NAME)
        sys.exit()


def find_max_column_table_data(data):
    min_column = 0

    if data is not None and len(data) > 0:
        for row in data:
            parts = row.split(EXCEL_DELIMITER)
            col_len = len(parts)
            if col_len > 0:
                if min_column == 0:
                    min_column = col_len - 1  # to remove last blank value added because of delimiter ==
                elif min_column > col_len - 1:
                    min_column = col_len = 1  # to remove last blank value added because of delimiter == -
    return min_column


def process_word_doc(doc, input_data, selection_data):
    if doc is not None and input_data is not None and len(input_data) > 0:
        global last_h1_updated
        global last_h2_updated
        global last_h3_updated
        global last_h4_updated
        global last_h5_updated
        global last_h6_updated
        para_style = PARA_STYLE
        # Table Data fetch with style
        table_data = []
        table_style = []
        for row in input_data:
            if row is not None and len(str(row).strip()) > 0:
                parts = str(row).strip().split(EXCEL_DELIMITER)
                if len(parts) >= 7:

                    col_h1 = parts[0].strip()
                    col_h2 = parts[1].strip()
                    col_h3 = parts[2].strip()
                    col_h4 = parts[3].strip()
                    col_h5 = parts[4].strip()
                    col_h6 = parts[5].strip()
                    col_content = parts[6].strip()
                    img_path = ""

                    if len(parts) >= 8 and len(parts[7].strip()) > 0:
                        para_style = parts[7].strip()

                    # Only check for table content in case all headings are same and content column is not having any data
                    if col_content is None or len(col_content) <= 0:
                        # Capture style for table only
                        if len(parts) >= 8 and len(parts[7].strip()) > 0:
                            table_style.append(parts[7].strip())

                        # Table Data Check
                        if len(parts) >= 9 and ((len(last_h1_updated) <= 0 and len(last_h2_updated) <= 0 and len(
                                last_h3_updated) <= 0 and len(last_h4_updated) <= 0 and len(
                            last_h5_updated) <= 0 and len(last_h5_updated) <= 0)
                                                or (
                                                        col_h1 == last_h1_updated and col_h2 == last_h2_updated and col_h3 == last_h3_updated and col_h4 == last_h4_updated and col_h5 ==
                                                        last_h5_updated and col_h6 == last_h6_updated)):
                            tr_data = ""
                            for data in parts[8:13]:
                                if len(data.strip()) > 0:
                                    tr_data = tr_data + data.strip() + EXCEL_DELIMITER
                            if len(str(tr_data)) > 0:
                                table_data.append(tr_data)
                            continue
                    else:
                        if len(parts) >= 13 and len(str(parts[13]).strip()) > 0:
                            img_path = str(parts[13]).strip()
                            print("Image Path: " + img_path)
                    max_table_column = 8
                    if table_data is not None and len(table_data) > 0:
                        max_table_column = find_max_column_table_data(table_data)
                        print("Table Data: " + str(table_data))
                        print("Table Max Column Data: " + str(max_table_column))
                        table = None
                        table = doc.add_table(rows=1, cols=max_table_column)
                        # For adding heading in table row
                        for row_table_index in range(0, len(table_data)):
                            row_parts = table_data[row_table_index].split(EXCEL_DELIMITER)

                            col_cells = table.rows[0].cells

                            for col_index in range(max_table_column):
                                col_cells[col_index].text = row_parts[col_index]
                            # To skip adding heading to rest of table rows
                            break

                        for row_table_index in range(1, len(table_data)):
                            row_parts = table_data[row_table_index].split(EXCEL_DELIMITER)
                            col_cells = table.add_row().cells
                            for col_index in range(max_table_column):
                                col_cells[col_index].text = row_parts[col_index]

                        if table_style is not None and len(table_style) > 0:
                            for style in table_style:
                                table.style = style
                                break

                        doc.add_paragraph('')
                    table_data = []
                    table_style = []

                    if (last_h1_updated is None or len(last_h1_updated.strip()) == 0 or not (
                            last_h1_updated == col_h1)) and len(col_h1) > 0 and col_h1 in selection_data[
                        KEY_H1_MAPPING]:
                        run = add_paragraph(doc, col_h1, H1_STYLE)
                        last_h1_updated = col_h1

                        last_h2_updated = last_h3_updated = last_h4_updated = last_h5_updated = last_h6_updated = ""

                    if col_h1 not in selection_data[KEY_H1_MAPPING]:
                        continue

                    if len(col_h2) <= 0 and col_content is not None and len(col_content) > 0:
                        add_paragraph(doc, col_content, para_style)

                        if img_path is not None and len(img_path) > 0:
                            add_paragraph(doc, "", "", img_path)

                        continue

                    if (last_h2_updated is None or len(last_h2_updated.strip()) == 0 or not (
                            last_h2_updated == col_h2)) and len(col_h2) > 0 and col_h2 in selection_data[
                        KEY_H2_MAPPING]:
                        run = add_paragraph(doc, col_h2, H2_STYLE)
                        last_h2_updated = col_h2

                        last_h3_updated = last_h4_updated = last_h5_updated = last_h6_updated = ""

                    if col_h2 not in selection_data[KEY_H2_MAPPING]:
                        continue

                    if len(col_h3) <= 0 and col_content is not None and len(col_content) > 0:
                        add_paragraph(doc, col_content, para_style)

                        if img_path is not None and len(img_path) > 0:
                            add_paragraph(doc, "", "", img_path)

                        continue

                    if (last_h3_updated is None or len(last_h3_updated.strip()) == 0 or not (
                            last_h3_updated == col_h3)) and len(col_h3) > 0 and col_h3 in selection_data[
                        KEY_H3_MAPPING]:
                        run = add_paragraph(doc, col_h3, H3_STYLE)
                        last_h3_updated = col_h3

                        last_h4_updated = last_h5_updated = last_h6_updated = ""

                    if col_h3 not in selection_data[KEY_H3_MAPPING]:
                        continue

                    if len(col_h4) <= 0 and col_content is not None and len(col_content) > 0:
                        add_paragraph(doc, col_content, para_style)

                        if img_path is not None and len(img_path) > 0:
                            add_paragraph(doc, "", "", img_path)

                        continue

                    if (last_h4_updated is None or len(last_h4_updated.strip()) == 0 or not (
                            last_h4_updated == col_h4)) and len(col_h4) > 0 and col_h4 in selection_data[
                        KEY_H4_MAPPING]:
                        run = add_paragraph(doc, col_h4, H4_STYLE)
                        last_h4_updated = col_h4

                        last_h5_updated = last_h6_updated = ""

                    if col_h4 not in selection_data[KEY_H4_MAPPING]:
                        continue

                    if len(col_h5) <= 0 and col_content is not None and len(col_content) > 0:
                        add_paragraph(doc, col_content, para_style)

                        if img_path is not None and len(img_path) > 0:
                            add_paragraph(doc, "", "", img_path)

                        continue

                    if (last_h5_updated is None or len(last_h5_updated.strip()) or not (
                            last_h5_updated == col_h5)) and len(col_h5) > 0 and col_h5 in selection_data[
                        KEY_H5_MAPPING]:
                        run = add_paragraph(doc, col_h5, H5_STYLE)
                        last_h5_updated = col_h5

                        last_h6_updated = ""

                    if col_h5 not in selection_data[KEY_H5_MAPPING]:
                        continue

                    if len(col_h6) <= 0 and col_content is not None and len(col_content) > 0:
                        add_paragraph(doc, col_content, para_style)

                        if img_path is not None and len(img_path) > 0:
                            add_paragraph(doc, "", "", img_path)

                        continue

                    if (last_h6_updated is None or len(last_h6_updated.strip()) == 0 or not (
                            last_h6_updated == col_h6)) and len(col_h6) > 0 and col_h6 in selection_data[
                        KEY_H6_MAPPING]:
                        run = add_paragraph(doc, col_h6, H6_STYLE)
                        last_h6_updated = col_h6
                        last_h6_updated = ""

                    if col_content is not None and len(col_content) > 0:
                        add_paragraph(doc, col_content)

                        if img_path is not None and len(img_path) > 0:
                            add_paragraph(doc, "", "", img_path)

                    print("Doc created")
        # Save the document
        doc.save(OUTPUT_FILE_NAME)
        print("Doc created: " + OUTPUT_FILE_NAME)
    else:
        print("Input data in blank. Aborting process!!!")


if __name__ == '__main__':
    start = time.time()
    input_data, selection_data, template_path = read_input_sheet_data()
    # Create a document
    try:
        doc = docx.Document(docx=template_path)
    except Exception as exp:
        print("File path to style template document could not be found, using default word document styles")
        doc = docx.Document()
    process_word_doc(doc, input_data, selection_data)
    end = time.time()
